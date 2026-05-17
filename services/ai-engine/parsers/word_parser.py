"""Stage 2: Word (.docx) parser"""
from typing import Optional
import io


def parse_word(file_content: bytes) -> str:
    """
    .docx fayldan matn chiqaradi.
    Qalin/kursiv formatni belgisi sifatida saqlaydi.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(file_content))
        lines = []

        for para in doc.paragraphs:
            if not para.text.strip():
                lines.append("")
                continue

            line_parts = []
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                # Qalin matn — to'g'ri javob belgisi bo'lishi mumkin
                if run.bold:
                    text = f"**{text}**"
                line_parts.append(text)

            lines.append("".join(line_parts))

        # Jadvallarni ham o'qish
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
            lines.append("")

        return "\n".join(lines)

    except Exception as e:
        raise RuntimeError(f"Word fayl o'qilmadi: {e}") from e
